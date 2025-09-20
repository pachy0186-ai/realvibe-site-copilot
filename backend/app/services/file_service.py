"""
File upload and processing service for RealVibe Site Copilot.
Handles document upload, text extraction, and storage.
"""

import os
import uuid
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
import mimetypes
from datetime import datetime

import PyPDF2
from docx import Document
import pandas as pd
from supabase import create_client, Client

from app.core.config import settings


class FileService:
    """Service for handling file uploads and processing."""
    
    def __init__(self):
        self.supabase: Client = create_client(
            settings.SUPABASE_URL,
            settings.SUPABASE_KEY
        )
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        
        # Supported file types
        self.supported_types = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'application/vnd.ms-excel': 'xls',
            'text/csv': 'csv'
        }
    
    async def upload_file(
        self, 
        file_content: bytes, 
        filename: str, 
        site_id: str,
        content_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Upload and process a file.
        
        Args:
            file_content: Raw file content
            filename: Original filename
            site_id: Site identifier
            content_type: MIME type of the file
            
        Returns:
            Dict containing file metadata and processing results
        """
        try:
            # Validate file type
            if not content_type:
                content_type, _ = mimetypes.guess_type(filename)
            
            if content_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Generate unique file ID and hash
            file_id = str(uuid.uuid4())
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Check for duplicate files
            existing_file = await self._check_duplicate(file_hash, site_id)
            if existing_file:
                return {
                    "file_id": existing_file["id"],
                    "message": "File already exists",
                    "duplicate": True
                }
            
            # Save file locally for processing
            file_extension = self.supported_types[content_type]
            local_path = self.upload_dir / f"{file_id}.{file_extension}"
            
            with open(local_path, 'wb') as f:
                f.write(file_content)
            
            # Extract text content
            text_content = await self._extract_text(local_path, content_type)
            
            # Upload to Supabase Storage
            storage_path = f"{site_id}/{file_id}.{file_extension}"
            storage_result = self.supabase.storage.from_("documents").upload(
                storage_path, file_content
            )
            
            if storage_result.get('error'):
                raise Exception(f"Storage upload failed: {storage_result['error']}")
            
            # Save file metadata to database
            file_metadata = {
                "id": file_id,
                "site_id": site_id,
                "filename": filename,
                "content_type": content_type,
                "file_size": len(file_content),
                "file_hash": file_hash,
                "storage_path": storage_path,
                "text_content": text_content,
                "page_count": self._get_page_count(local_path, content_type),
                "upload_time": datetime.utcnow().isoformat(),
                "processing_status": "completed"
            }
            
            db_result = self.supabase.table("files").insert(file_metadata).execute()
            
            if db_result.data:
                # Clean up local file
                local_path.unlink()
                
                return {
                    "file_id": file_id,
                    "filename": filename,
                    "file_size": len(file_content),
                    "page_count": file_metadata["page_count"],
                    "text_length": len(text_content),
                    "storage_path": storage_path,
                    "duplicate": False
                }
            else:
                raise Exception("Failed to save file metadata to database")
                
        except Exception as e:
            # Clean up on error
            if 'local_path' in locals() and local_path.exists():
                local_path.unlink()
            
            raise Exception(f"File upload failed: {str(e)}")
    
    async def _extract_text(self, file_path: Path, content_type: str) -> str:
        """Extract text content from various file types."""
        try:
            if content_type == 'application/pdf':
                return self._extract_pdf_text(file_path)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return self._extract_docx_text(file_path)
            elif content_type == 'text/plain':
                return self._extract_txt_text(file_path)
            elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'text/csv']:
                return self._extract_spreadsheet_text(file_path)
            else:
                return ""
        except Exception as e:
            print(f"Text extraction failed for {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PDF text extraction failed: {e}")
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"DOCX text extraction failed: {e}")
        return text.strip()
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"TXT text extraction failed: {e}")
            return ""
    
    def _extract_spreadsheet_text(self, file_path: Path) -> str:
        """Extract text from spreadsheet files."""
        text = ""
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert dataframe to text representation
            text = df.to_string(index=False)
        except Exception as e:
            print(f"Spreadsheet text extraction failed: {e}")
        return text
    
    def _get_page_count(self, file_path: Path, content_type: str) -> int:
        """Get page count for supported file types."""
        try:
            if content_type == 'application/pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    return len(pdf_reader.pages)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                # Approximate page count for DOCX (rough estimate)
                doc = Document(file_path)
                return max(1, len(doc.paragraphs) // 50)  # Rough estimate
            else:
                return 1
        except Exception:
            return 1
    
    async def _check_duplicate(self, file_hash: str, site_id: str) -> Optional[Dict]:
        """Check if file with same hash already exists for the site."""
        try:
            result = self.supabase.table("files").select("*").eq(
                "file_hash", file_hash
            ).eq("site_id", site_id).execute()
            
            return result.data[0] if result.data else None
        except Exception:
            return None
    
    async def get_file_list(self, site_id: str) -> List[Dict[str, Any]]:
        """Get list of files for a site."""
        try:
            result = self.supabase.table("files").select(
                "id, filename, content_type, file_size, page_count, upload_time, processing_status"
            ).eq("site_id", site_id).order("upload_time", desc=True).execute()
            
            return result.data or []
        except Exception as e:
            print(f"Failed to get file list: {e}")
            return []
    
    async def delete_file(self, file_id: str, site_id: str) -> bool:
        """Delete a file and its associated data."""
        try:
            # Get file metadata
            file_result = self.supabase.table("files").select("*").eq(
                "id", file_id
            ).eq("site_id", site_id).execute()
            
            if not file_result.data:
                return False
            
            file_data = file_result.data[0]
            
            # Delete from storage
            self.supabase.storage.from_("documents").remove([file_data["storage_path"]])
            
            # Delete from database
            self.supabase.table("files").delete().eq("id", file_id).execute()
            
            return True
        except Exception as e:
            print(f"Failed to delete file: {e}")
            return False
    
    async def get_file_content(self, file_id: str, site_id: str) -> Optional[Dict[str, Any]]:
        """Get file content and metadata."""
        try:
            result = self.supabase.table("files").select("*").eq(
                "id", file_id
            ).eq("site_id", site_id).execute()
            
            return result.data[0] if result.data else None
        except Exception as e:
            print(f"Failed to get file content: {e}")
            return None

